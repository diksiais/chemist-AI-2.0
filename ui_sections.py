# ui_sections.py
import streamlit as st
from urllib.parse import quote # For encoding URLs in chemical lookup
import re # For parsing ideas
import io # For handling file downloads (BytesIO)
from docx import Document # New import for DOCX generation

# Import functions from other modules that UI sections need
from workflow import (
    generate_research_ideas_from_ai,
    refine_single_idea_from_ai,
    answer_follow_up_question_from_ai,
    suggest_search_queries_from_ai,
    generate_literature_summary_from_ai,
    generate_properties_from_ai,
    compile_final_response_from_ai,
    perform_chemical_lookup
)
from database import load_search_history, save_search_history, delete_search_history_entry, clear_all_search_history
from pdf_processor import extract_text_from_pdf, get_combined_uploaded_text


def render_input_details_stage():
    """Renders the UI for Step 1: Provide Research Details."""
    st.subheader("Step 1: Provide Research Details")

    # Load history data for display
    st.session_state.search_history_data = load_search_history()

    # History selection dropdown
    history_options = ["--- Select from History ---"] + [
        f"{entry['timestamp']} - {entry['topic']}" for entry in st.session_state.search_history_data
    ]
    selected_option = st.selectbox(
        "Load Previous Search:",
        options=history_options,
        index=0,
        key="history_selector"
    )

    # Populate inputs if a history item is selected
    selected_entry = None
    if selected_option != "--- Select from History ---":
        selected_index = history_options.index(selected_option) - 1
        selected_entry = st.session_state.search_history_data[selected_index]
        st.session_state.selected_history_id = selected_entry['id']
    else:
        st.session_state.selected_history_id = None

    # Input fields
    topic = st.text_input(
        "🔍 Research Topic",
        value=selected_entry['topic'] if selected_entry else st.session_state.current_topic, # Retain current input if not from history
        placeholder="e.g., Sustainable catalysts for plastic degradation",
        key="input_topic"
    )
    goal = st.text_area(
        "🎯 Research Goal",
        value=selected_entry['goal'] if selected_entry else st.session_state.current_goal,
        placeholder="e.g., Develop a highly efficient and reusable catalyst that can break down PET plastics into their monomers at room temperature.",
        key="input_goal"
    )
    data = st.text_area(
        "🧪 Data We Already Have",
        value=selected_entry['data'] if selected_entry else st.session_state.current_data,
        placeholder="e.g., Initial screening results of metal-organic frameworks (MOFs) showing some catalytic activity, spectroscopic data of degraded plastic samples.",
        key="input_data"
    )

    st.markdown("---")
    st.subheader("Upload Research Papers (Optional)")
    uploaded_files = st.file_uploader(
        "Upload PDF research papers for AI to use as reference:",
        type=["pdf"],
        accept_multiple_files=True,
        key="pdf_uploader"
    )

    # Process newly uploaded files
    if uploaded_files:
        current_uploaded_names = {p['name'] for p in st.session_state.uploaded_papers_data}
        new_files_uploaded = False # Flag to check if any new file was processed
        for uploaded_file in uploaded_files:
            if uploaded_file.name not in current_uploaded_names:
                with st.spinner(f"Extracting text from {uploaded_file.name}..."):
                    extracted_text = extract_text_from_pdf(uploaded_file)
                    if not "Error extracting text" in extracted_text:
                        st.session_state.uploaded_papers_data.append({
                            'name': uploaded_file.name,
                            'extracted_text': extracted_text
                        })
                        st.success(f"Successfully extracted text from '{uploaded_file.name}'.")
                        new_files_uploaded = True
                    else:
                        st.error(f"Failed to extract text from '{uploaded_file.name}'. Please try another file.")
        # Only rerun if new files were actually added to avoid excessive reruns
        if new_files_uploaded:
            st.rerun() # Keep rerun here to update the displayed list of uploaded papers immediately

    # Display list of currently uploaded papers
    if st.session_state.uploaded_papers_data:
        st.markdown("**Currently Uploaded Papers:**")
        for i, paper_info in enumerate(st.session_state.uploaded_papers_data):
            st.write(f"- {paper_info['name']}")
        
        if st.button("Clear All Uploaded Papers", key="clear_uploaded_papers"):
            st.session_state.uploaded_papers_data = []
            st.success("All uploaded papers cleared.")
            st.rerun()
    st.markdown("---")


    col_buttons = st.columns(2)
    with col_buttons[0]:
        if st.button("💡 Generate Research Ideas", help="Click to generate research ideas based on your input."):
            if not topic or not goal or not data:
                st.warning("Please fill in all fields to generate research ideas.")
            else:
                with st.spinner("Generating ideas... This might take a moment."):
                    # Store current inputs in session state for refinement context
                    st.session_state.current_topic = topic
                    st.session_state.current_goal = goal
                    st.session_state.current_data = data
                    save_search_history(topic, goal, data)
                    st.session_state.ideas = generate_research_ideas_from_ai(topic, goal, data)
                    st.session_state.idea_index = 0
                    if st.session_state.ideas:
                        st.session_state.stage = 'review_ideas'
                    st.rerun()
    with col_buttons[1]:
        if st.session_state.search_history_data:
            if st.button("🗑️ Clear All History", help="Delete all saved search history entries."):
                clear_all_search_history()
                st.success("All search history cleared!")
                st.rerun()

    if st.session_state.search_history_data:
        st.markdown("---")
        st.subheader("Manage Individual History Entries")
        for entry in st.session_state.search_history_data:
            col_entry_display, col_entry_delete = st.columns([0.8, 0.2])
            with col_entry_display:
                st.markdown(f"**{entry['timestamp']}** - {entry['topic']}")
                with st.expander("Details"):
                    st.write(f"**Goal:** {entry['goal']}")
                    st.write(f"**Data:** {entry['data']}")
            with col_entry_delete:
                if st.button("Delete", key=f"delete_history_{entry['id']}"):
                    delete_search_history_entry(entry['id'])
                    st.success(f"Entry '{entry['topic']}' deleted.")
                    st.rerun()


def render_review_ideas_stage():
    """Renders the UI for Step 2: Review Research Idea."""
    st.subheader(f"Step 2: Review Research Idea #{st.session_state.idea_index + 1}")

    if st.session_state.ideas and st.session_state.idea_index < len(st.session_state.ideas):
        idea = st.session_state.ideas[st.session_state.idea_index]
        st.markdown(f"**Idea:** {idea}")

        st.markdown("---")
        st.subheader("Refine This Idea?")
        refinement_feedback = st.text_area(
            "Provide feedback to refine this idea (e.g., 'Make it more specific to polymers', 'Suggest alternative synthesis methods'):",
            key="refinement_feedback"
        )
        if st.button("🔄 Refine Current Idea"):
            if refinement_feedback.strip():
                with st.spinner("Refining idea..."):
                    refined_idea = refine_single_idea_from_ai(
                        original_idea=idea,
                        refinement_feedback=refinement_feedback,
                        topic=st.session_state.current_topic,
                        goal=st.session_state.current_goal,
                        data=st.session_state.current_data
                    )
                    st.session_state.ideas[st.session_state.idea_index] = refined_idea
                    st.success("Idea refined!")
                    st.rerun()
            else:
                st.warning("Please enter feedback to refine the idea.")
        st.markdown("---")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("👍 Approve Idea"):
                st.session_state.approved_idea = idea
                st.session_state.stage = 'literature_summary'
                st.rerun()
        with col2:
            if st.button("👎 Disapprove & Next Idea"):
                st.session_state.idea_index += 1
                if st.session_state.idea_index >= len(st.session_state.ideas):
                    st.warning("No more ideas to review. Please go back to generate new ideas.")
                    st.session_state.stage = 'input_details'
                st.rerun()
    else:
        st.warning("No ideas available. Please generate new ideas.")
        st.session_state.stage = 'input_details'


def render_literature_summary_stage():
    """Renders the UI for Step 3: Generate Literature Summary."""
    st.subheader("Step 3: Generate Literature Summary")
    st.markdown(f"**Approved Idea:** {st.session_state.approved_idea}")

    if st.session_state.literature_summary is None:
        with st.spinner("Generating literature summary..."):
            st.session_state.literature_summary = generate_literature_summary_from_ai(st.session_state.approved_idea)

    st.markdown("---")
    st.markdown("**Generated Literature Summary:**")
    if st.session_state.literature_summary and "⚠️ Error:" not in st.session_state.literature_summary:
        st.text_area(
            "Copy Literature Summary below:",
            value=st.session_state.literature_summary,
            height=300,
            key="literature_summary_output"
        )
        # Download as DOCX button
        doc = Document()
        doc.add_heading("Literature Summary", level=1)
        doc.add_paragraph(st.session_state.literature_summary)
        
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        st.download_button(
            label="Download Summary as DOCX",
            data=bio.getvalue(),
            file_name="literature_summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_literature_docx"
        )
    else:
        st.error(st.session_state.literature_summary if st.session_state.literature_summary else "No summary generated.")
    st.markdown("---")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("👍 Approve Summary"):
            st.session_state.stage = 'properties_prediction'
            # Reset follow-up question/response and search queries when moving to next stage
            st.session_state.follow_up_question = ""
            st.session_state.follow_up_response = None
            st.session_state.suggested_search_queries = [] # Reset
            st.rerun()
    with col2:
        if st.button("👎 Disapprove & Re-evaluate Idea"):
            st.session_state.literature_summary = None
            st.session_state.stage = 'review_ideas'
            st.session_state.idea_index += 1
            if st.session_state.idea_index >= len(st.session_state.ideas):
                st.warning("No more ideas to review. Please go back to generate new ideas.")
                st.session_state.stage = 'input_details'
            # Reset follow-up question/response and search queries
            st.session_state.follow_up_question = ""
            st.session_state.follow_up_response = None
            st.session_state.suggested_search_queries = [] # Reset
            st.rerun()

    # --- AI-Driven Follow-up Questions Section ---
    st.markdown("---")
    st.subheader("Ask a Follow-up Question (AI)")
    st.session_state.follow_up_question = st.text_input(
        "Your question:",
        value=st.session_state.follow_up_question,
        key="follow_up_question_input_lit_summary"
    )
    if st.button("Ask AI", key="ask_ai_lit_summary_button"):
        if st.session_state.follow_up_question.strip():
            with st.spinner("Getting AI's answer..."):
                st.session_state.follow_up_response = answer_follow_up_question_from_ai(
                    st.session_state.approved_idea,
                    st.session_state.literature_summary,
                    st.session_state.properties, # Pass properties even if not yet generated, it will be None
                    st.session_state.follow_up_question
                )
            st.rerun() # Rerun to display the response
        else:
            st.warning("Please enter a question.")

    if st.session_state.follow_up_response:
        st.markdown("**AI's Answer:**")
        st.info(st.session_state.follow_up_response)

    # --- AI-Suggested Search Queries Section ---
    st.markdown("---")
    st.subheader("AI-Suggested Search Queries")
    if st.button("Suggest Search Queries", key="suggest_queries_button"):
        with st.spinner("Generating search queries..."):
            st.session_state.suggested_search_queries = suggest_search_queries_from_ai(
                st.session_state.approved_idea,
                st.session_state.literature_summary
            )
        st.rerun() # Rerun to display the queries

    if st.session_state.suggested_search_queries:
        st.markdown("Here are some queries you might find useful:")
        for i, query in enumerate(st.session_state.suggested_search_queries):
            # Encode query for URL
            encoded_query = quote(query)
            pubmed_url = f"https://pubmed.ncbi.nlm.nih.gov/?term={encoded_query}"
            scopus_url = f"https://www.scopus.com/results/results.uri?sort=plf-f&src=s&st1={encoded_query}&sid=0a0123456789abcdef0123456789abcdef&sot=b&sdt=b&sl=0&s=TITLE-ABS-KEY%28{encoded_query}%29&origin=resultslist&zone=resultsList"
            google_scholar_url = f"https://scholar.google.com/scholar?q={encoded_query}"

            st.markdown(f"- **{query}** ([PubMed]({pubmed_url}) | [Scopus]({scopus_url}) | [Google Scholar]({google_scholar_url}))")


def render_properties_prediction_stage():
    """Renders the UI for Step 4: Predict Properties / Experimental Approach."""
    st.subheader("Step 4: Predict Properties / Experimental Approach")
    st.markdown(f"**Approved Idea:** {st.session_state.approved_idea}")

    if st.session_state.properties is None:
        with st.spinner("Generating property predictions..."):
            st.session_state.properties = generate_properties_from_ai(st.session_state.approved_idea)

    st.markdown("---")
    st.markdown("**Generated Properties/Approach (AI):**")
    st.markdown(st.session_state.properties)
    st.markdown("---")

    st.subheader("Chemical Structure Lookup")
    st.session_state.chemical_query_input = st.text_input(
        "Enter Chemical Name or CAS Number:",
        value=st.session_state.chemical_query_input,
        key="chemical_lookup_input"
    )
    
    # Store the state of the lookup attempt
    if 'chemical_lookup_attempted' not in st.session_state:
        st.session_state.chemical_lookup_attempted = False
    if 'chemical_lookup_success' not in st.session_state:
        st.session_state.chemical_lookup_success = False

    if st.button("🔎 Look Up Chemical Structure"):
        st.session_state.chemical_lookup_attempted = True # Mark that an attempt was made
        st.session_state.chemical_lookup_success = False # Assume failure until proven otherwise
        
        # Clear previous lookup results before a new attempt
        st.session_state.chemical_cid = None
        st.session_state.chemical_image_url = None
        st.session_state.chemical_source = None
        st.session_state.chemical_matched_name = None
        st.session_state.chemical_image_bytes = None

        if st.session_state.chemical_query_input:
            with st.spinner(f"Looking up '{st.session_state.chemical_query_input}'..."):
                cid, image_url, source, matched_name, image_bytes = perform_chemical_lookup(st.session_state.chemical_query_input)
                
                st.session_state.chemical_cid = cid
                st.session_state.chemical_image_url = image_url
                st.session_state.chemical_source = source
                st.session_state.chemical_matched_name = matched_name
                st.session_state.chemical_image_bytes = image_bytes

                if image_url:
                    st.session_state.chemical_lookup_success = True
                # No st.rerun() here, allow the rest of the script to execute and display messages
        else:
            st.warning("Please enter a chemical name or CAS number to look up.")
            st.session_state.chemical_lookup_attempted = False # No valid input, so no real attempt

    # Display chemical lookup results or error/suggestion messages
    if st.session_state.chemical_lookup_attempted:
        if st.session_state.chemical_lookup_success:
            st.markdown("---")
            st.markdown("**Found Chemical Structure:**")
            st.image(st.session_state.chemical_image_url, caption=f"{st.session_state.chemical_matched_name} (Source: {st.session_state.chemical_source})", use_column_width=True)
            if st.session_state.chemical_cid:
                st.info(f"PubChem CID: {st.session_state.chemical_cid}")
            
            # Download image button
            if st.session_state.chemical_image_bytes:
                st.download_button(
                    label="Download Structure Image",
                    data=st.session_state.chemical_image_bytes,
                    file_name=f"{st.session_state.chemical_matched_name}_structure.png",
                    mime="image/png",
                    key="download_structure_image"
                )
            st.markdown("---")
        else:
            # Only show this warning/info if a lookup was attempted and failed
            st.warning("🔍 Compound not found in structured chemical databases.")
            query_encoded = quote(st.session_state.chemical_query_input)
            pubchem_url = f"https://pubchem.ncbi.nlm.nih.gov/#query={encoded_query}"
            wikidata_url = f"https://www.wikidata.org/w/index.php?search={encoded_query}"

            st.info(
                "⚠️ This compound may still exist in scientific literature, but no structure or CID was found.\n\n"
                "💡 Try entering a more precise identifier such as:\n"
                "- CAS number (e.g. `50-00-0`)\n"
                "- IUPAC name (e.g. `methanal`)\n"
                "- SMILES notation (e.g. `C=O`)\n\n"
                f"You can also try searching manually:\n"
                f"- 🔬 [Search on PubChem]({pubchem_url})\n"
                f"- 🧠 [Search on Wikidata]({wikidata_url})"
            )
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("👍 Approve Properties"):
            st.session_state.stage = 'final_compilation'
            # Reset follow-up question/response and search queries when moving to next stage
            st.session_state.follow_up_question = ""
            st.session_state.follow_up_response = None
            st.session_state.suggested_search_queries = [] # Reset
            st.rerun()
    with col2:
        if st.button("👎 Disapprove & Re-evaluate Summary"):
            st.session_state.properties = None
            st.session_state.literature_summary = None
            st.session_state.stage = 'literature_summary'
            # Reset follow-up question/response and search queries
            st.session_state.follow_up_question = ""
            st.session_state.follow_up_response = None
            st.session_state.suggested_search_queries = [] # Reset
            st.rerun()

    # --- AI-Driven Follow-up Questions Section ---
    st.markdown("---")
    st.subheader("Ask a Follow-up Question (AI)")
    st.session_state.follow_up_question = st.text_input(
        "Your question:",
        value=st.session_state.follow_up_question,
        key="follow_up_question_input_props_pred"
    )
    if st.button("Ask AI", key="ask_ai_props_pred_button"):
        if st.session_state.follow_up_question.strip():
            with st.spinner("Getting AI's answer..."):
                st.session_state.follow_up_response = answer_follow_up_question_from_ai(
                    st.session_state.approved_idea,
                    st.session_state.literature_summary,
                    st.session_state.properties,
                    st.session_state.follow_up_question
                )
            st.rerun() # Rerun to display the response
        else:
            st.warning("Please enter a question.")

    if st.session_state.follow_up_response:
        st.markdown("**AI's Answer:**")
        st.info(st.session_state.follow_up_response)

    # --- AI-Suggested Search Queries Section ---
    st.markdown("---")
    st.subheader("AI-Suggested Search Queries")
    if st.button("Suggest Search Queries", key="suggest_queries_button_props_pred"):
        with st.spinner("Generating search queries..."):
            st.session_state.suggested_search_queries = suggest_search_queries_from_ai(
                st.session_state.approved_idea,
                st.session_state.literature_summary # Pass literature summary even if it's from previous stage
            )
        st.rerun() # Rerun to display the queries

    if st.session_state.suggested_search_queries:
        st.markdown("Here are some queries you might find useful:")
        for i, query in enumerate(st.session_state.suggested_search_queries):
            # Encode query for URL
            encoded_query = quote(query)
            pubmed_url = f"https://pubmed.ncbi.nlm.nih.gov/?term={encoded_query}"
            scopus_url = f"https://www.scopus.com/results/results.uri?sort=plf-f&src=s&st1={encoded_query}&sid=0a0123456789abcdef0123456789abcdef&sot=b&sdt=b&sl=0&s=TITLE-ABS-KEY%28{encoded_query}%29&origin=resultslist&zone=resultsList"
            google_scholar_url = f"https://scholar.google.com/scholar?q={encoded_query}"

            st.markdown(f"- **{query}** ([PubMed]({pubmed_url}) | [Scopus]({scopus_url}) | [Google Scholar]({google_scholar_url}))")


def render_final_compilation_stage():
    """Renders the UI for Step 5: Final Research Proposal Overview."""
    st.subheader("Step 5: Final Research Proposal Overview")
    st.markdown(f"**Approved Idea:** {st.session_state.approved_idea}")
    st.markdown(f"**Literature Summary:** {st.session_state.literature_summary}")
    st.markdown(f"**Properties/Approach:** {st.session_state.properties}")
    if st.session_state.chemical_image_url:
        st.markdown(f"**Looked Up Chemical:** {st.session_state.chemical_matched_name} (Source: {st.session_state.chemical_source})")
        st.image(st.session_state.chemical_image_url, caption=f"Structure of {st.session_state.chemical_matched_name}", use_column_width=True)
        # Download image button in final stage too
        if st.session_state.chemical_image_bytes:
            st.download_button(
                label="Download Structure Image (Final)",
                data=st.session_state.chemical_image_bytes,
                file_name=f"{st.session_state.chemical_matched_name}_structure_final.png",
                mime="image/png",
                key="download_structure_image_final"
            )


    if st.session_state.final_response is None:
        with st.spinner("Compiling final response..."):
            st.session_state.final_response = compile_final_response_from_ai(
                st.session_state.approved_idea,
                st.session_state.literature_summary,
                st.session_state.properties
            )

    st.markdown("---")
    st.markdown("**Final Research Proposal Overview:**")
    full_content_to_copy = (
        f"Research Idea: {st.session_state.approved_idea}\n\n"
        f"Literature Summary:\n{st.session_state.literature_summary}\n\n"
        f"Properties/Approach:\n{st.session_state.properties}\n\n"
    )
    if st.session_state.chemical_matched_name:
        full_content_to_copy += f"Chemical Looked Up: {st.session_state.chemical_matched_name} (Source: {st.session_state.chemical_source})\n\n"
    full_content_to_copy += f"Final Proposal Overview:\n{st.session_state.final_response}"

    if st.session_state.final_response and "⚠️ Error:" not in st.session_state.final_response:
        st.text_area(
            "Copy Final Research Proposal below:",
            value=full_content_to_copy,
            height=500,
            key="final_proposal_output"
        )
        # Download as DOCX button for final proposal
        doc = Document()
        doc.add_heading("Final Research Proposal Overview", level=1)
        # Add the full content to the document, splitting by newlines for better formatting
        for paragraph_text in full_content_to_copy.split('\n'):
            doc.add_paragraph(paragraph_text)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        st.download_button(
            label="Download Proposal as DOCX",
            data=bio.getvalue(),
            file_name="research_proposal.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_proposal_docx"
        )
    else:
        st.error(st.session_state.final_response if st.session_state.final_response else "No final proposal generated.")
    st.markdown("---")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Start New Research"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()
